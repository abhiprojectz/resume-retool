from docx import Document
from bs4 import BeautifulSoup
from docx.shared import Pt, RGBColor

def convert_html_to_docx(html_file, docx_file):
    # Load the HTML file
    with open(html_file, 'r') as file:
        html_content = file.read()

    # Parse the HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # Create a new Word document
    doc = Document()

    # Iterate over each element in the HTML file
    for element in soup.descendants:
        if element.name:
        # Check the type of element
            if element.name == 'h1':
                # Heading element
                level = int(element.name[1])
                paragraph = doc.add_paragraph(element.get_text(), style='Title')
                paragraph.style.font.color.rgb = RGBColor(0, 0, 0)  

            elif element.name == 'p':
                # Paragraph element
                paragraph = doc.add_paragraph(element.get_text())
                # Change the font size to 12
                paragraph.style.font.size = Pt(12)
                # Add a long underline
                paragraph.runs[0].underline = True
                # paragraph.runs[0].underline_color.rgb = RGBColor(0, 0, 0)  # Black color
                # Change the color of the paragraph to red
                paragraph.style.font.color.rgb = RGBColor(255, 0, 0)  # Red color

            elif element.name == 'ul':
                bullet_list = doc.add_paragraph()

                # Iterate over each list item (li) in the unordered list
                for li in element.find_all('li'):
                    # Add the list item as a bullet point in the Word document
                    bullet_list.add_run(li.get_text()).bold = True
                    bullet_list.add_run('\n')

    # Save the Word document
    doc.save(docx_file)




convert_html_to_docx('page.html', 'output.docx')