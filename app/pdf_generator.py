import markdown
import json
import os
import requests
# import pdfkit
import PyPDF2
import textwrap
import docx


EDEN_AI_API_KEY = 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ1c2VyX2lkIjoiZDA0Yjk5OTYtODMwYi00YWVjLTllYzgtOGQ1MmExYTk5MGZkIiwidHlwZSI6ImFwaV90b2tlbiJ9.ORxx6jIpcCzK95cZ-yVpSe_tOEP0QAjoq92MlfDqixE'
CLIENT_ID = "71d2c5ded07748889b86e68730b5780e"
CLIENT_SECRET = "p8e-7zzsRzTT1Bb9XK1ASvKFsn_5uhETb5E9"



def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)


def extract_text_from_pdf(pdf_path):
    text = ""
    with open(pdf_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        num_pages = len(reader.pages)

        for page_num in range(num_pages):
            page = reader.pages[page_num]
            text += page.extract_text()
    return text


def generate_gpt(text, history=[]):
        headers = {"Authorization": f"Bearer {EDEN_AI_API_KEY}"}
        url ="https://api.edenai.run/v2/text/chat"
        payload = {
            "providers": "openai",
            "text": text,
            "chat_global_action": "Follow user instructions",
            "previous_history" : history,
            "temperature" : 0.0,
            "settings":{"openai":"gpt-3.5-turbo"},
            "max_tokens" : 1000
            }
        response = requests.post(url, json=payload, headers=headers)
        # print(response.text)
        try:
            result = json.loads(response.text)
            msg = result['openai']['generated_text']
            # print(msg)
        except Exception as e:
            return
        return msg


def extract_links_from_pdf(file_path):
    PDFFile = open(file_path, 'rb')
    PDF = PyPDF2.PdfReader(PDFFile)
    pages = len(PDF.pages)
    key = '/Annots'
    uri = '/URI'
    ank = '/A'
    links = []

    for page in range(pages):
        pageSliced = PDF.pages[page]
        pageObject = pageSliced.get_object()
        if key in pageObject.keys():
            ann = pageObject[key]
            for a in ann:
                u = a.get_object()
                if uri in u[ank].keys():
                    links.append(u[ank][uri])

    return '\n'.join(links)


PROMPT = '''
You are given a raw text from a Resume, Your task is to write a clean well formatted extracted information by analyzing the given text. Hints/help on what to do with are provided & enclosed by [] brackets. Your output should only be in well formatted markdown. No explanations or replies or Note text required or any information Just need the output Omit extra details out of resume format.

Here's the format:

# Person's name here
[Extract person's information and output in below format, IGNORE IF some info is not provided]
[City], [State]  |  [Phone Number]  |  [Email Address]  |  [other LINKS, sperated by |]


---

## SUMMARY
Write a short summary of overall text (e.g., project manager, data analyst) with a X-year track record of driving results. Career Highlight 1 (What you did -> What the result was). Career Highlight 2 (e.g., Led development of new $30 million aircraft component that shipped 3 months early). Career Highlight 3 (e.g., Renegotiated supply contract with equipment vendor, saving $500K in the first year)

## EXPERIENCE
[Limit to only first 2 experiences, ignore REST]

**Translated Job Title** (e.g., IT Specialist vs. Signal Systems Specialist), Branch of Military, Location
<span>05/2017 - 01/2019</span>


Write a 1-line overview of the role that describes your scope of responsibilities
Include 3-5 results-oriented bullet points per section
Bullet point should only be 1-2 lines long
Add measurable metrics and results to your bullets whenever possible

Example of this: **Project Manager**, US Army, Ft. Benning, Georgia
<span>05/2017 - 01/2019</span>


Led 20-person team responsible for training 800 new personnel with an annual budget of $1M+.
Developed new physical training program with clear goals and checkpoints, improving average scores 85%.
Took the lowest-performing training team to the top-ranked one by earning buy-in early on and then holding team members accountable for the performance of their trainees.
Oversaw forecasting for all necessary supplies; identified economies of scale that led to a 27% reduction in costs and 100% availability of required resources.

## EDUCATION
[Limit to only first 3 education, ignore REST]

**Write here the extracted Title of education**
<span>05/2017 - 01/2019</span>
(Write the Program/Specialty)
Write here the extracted University/Institution Name, Location, Honors


## ADDITIONAL INFORMATION
[Keep short & concise, ignore REST]
Write a succinct description of relevant certifications, skills, languages, awards, etc. [List in order of relevancy to job application.]
+ **Technical Skills**: (Just output level of proficiency) i.e. SQL (advanced), HTML (proficient), etc.
+ **Certifications**: i.e. Black Belt in Six Sigma, PMP
+ **Awards**: i.e. Officer of the Year (2014), ranked #1 of 80 for enhancements to operational efficiency
+ **Languages**: (Just output level of proficiency), i.e. French (advanced), Spanish (basic), German (conversational)
+ **Interests**: (include examples), i.e. running (completed 4 marathons across 3 continents), cooking (won Ft. Benning’s annual Chili Cook-Off)


Here's the Raw text:
'''

_BASE_HELPER_PROMPT = '''
You are given a Portion of a raw text from a Resume, Your task is to just shorten the given portion of the text and output required information only that should be a part of the final format.
write a clean well formatted extracted information by analyzing the given text. Hints/help on what to do with are provided & enclosed by [] brackets. No explanations or replies or Note text required or any information Just need the output Omit extra details out of resume format. Note that the raw text does not have all portion as its just a small portion of it, You only need to minimize and clarify that portion so it could be later used in final version.
Just limit the portion to one or two blocks. OFCOURCE SOME PORTION MAY NOT BE THERE JUST IGNORE IT , SHORT & CONCISE OUTPUT REQUIRED.
Include every possible information like timeline worked for role or education etc only if present else ignore or some additional information.
Format:

Person's name here [If present]
[Extract person's information and output in below format, IGNORE IF some info is not provided]
[City], [State] ǀ [Phone Number] ǀ [Email Address] ǀ [other LINKS, sperated by | (pipe character)]

Ignore if there's none.

Here's the Raw text:
'''

CLEAN_PROMPT = '''
Act as a resume cleaner tool, You are given an resume text in markdown format. You need to omit remove the Extra fields and section that are not listed in below format. Do not write or explain anything. Do not write any greetings text, i need only pure cleaned resume text output.
Format:

## SUMMARY
[Including Its inner content ONLY]
## EXPERIENCE
[Including Its inner content ONLY]
## EDUCATION
[Including Its inner content ONLY]
## ADDITIONAL INFORMATION
[Including Its inner content ONLY]

Except only the above 5 section, if any extra section or text is present just remove it from the text. The output should only have these 5 sections please DO NOT ADD ANY EXTRA SECTION, REMOVE any extra section it is important to have only this format. DO NOT CHANGE ANY FORMAT OR ADD ANYTHING.

'''




BASE_PROMPT = '''
You are given a protion of a candidate's resume raw text, and you need to just output the asked information in the given desired format only. Do not write or explain anything except what is being asked. Do not write any replies or greetings message. Only output in desired MARKDOWN format. Do not write anything extra information except the format. Simply ignore any extra information that is given just output what is being asked. Make sure your output matches the format and only contains the required data. The hints or asked information is enclosed by [] brackets. DO NOT WRITE THEM in you routput it is for your understanding only, use for that purpose only.

[DO NOT ADD ANY EXTRA SECTION OR HEADINGS OR SUB HEADINGS EXCEPT THE ONE SPECIFIED in ABOVE FORMAT. ONLY OUTPUT ABOVE SECTION NO EXTRA CONTENT OR SECTION REQUIRED. FOLLOW THE FORMAT STRICTLY]
Format:
'''

INFO_SECTION = '''
# Write the CANDIDATE'S ACTUAL NAME HERE 
[Extract CANDIDATE's information and only output in below format. Do not write anything except the below asked information. Add the relevant links in thier accrurate position. IGNORE the rest of the IF some info is not provided, Do not write the hints if info is not provided to it just fully ignore that like if email address or other links ETC not provided do not write thier blank format, ONLY WRITE ACTUAL DATA , IGNORE IF NOT PROVIDED.]
[City], [State]  |  [Phone Number]  |  [Email Address]  |  [other LINKS, sperated by |]

[DO NOT ADD ANY EXTRA SECTION OR HEADINGS EXCEPT THE ONE SPECIFIED ABOVE FORMAT.]
Resume text:
'''

SUMMARY_SECTION = '''
## SUMMARY
Write a short summary (THE SHORT SUMMARY SHOULD NOT EXCEED 3 LINES) of overall text (e.g., project manager, data analyst) with a X-year track record of driving results. Career Highlight 1 (What you did -> What the result was). Career Highlight 2 (e.g., Led development of new $30 million aircraft component that shipped 3 months early). Career Highlight 3 (e.g., Renegotiated supply contract with equipment vendor, saving $500K in the first year) 

[DO NOT ADD ANY EXTRA SECTION OR HEADINGS OR SUB HEADINGS EXCEPT THE ONE SPECIFIED in ABOVE FORMAT. ONLY OUTPUT ABOVE SECTION NO EXTRA CONTENT OR SECTION REQUIRED. FOLLOW THE FORMAT STRICTLY]
Resume text:
'''

EXPERIENCE_SECTION = '''
## EXPERIENCE
[OUTPUT ONLY THE FIRST 2 EXPERIENCE OUT OF THE TEXT IGNORE ALL REST, FOLLOW THE experiences format specifically, write in short. Do not write more than 2 experience, write only first 2. Thier content should follow below format]

**Translated Job Title** (e.g., IT Specialist vs. Signal Systems Specialist), Branch of Military, Location  
[NOW LEAVE A LINE AND PASTE THE DURATION ENCLOSED WITH THE <span> tag, Strictly follow this, DO NOT WRITE DURATION IN PLAIN TEXT]
<span>Jan 2001 - Dec 2002</span> [WRITE THE DURATION PERIOD ENCLOSED WITH <span> tag]

Write a 1-line overview of the role that describes your scope of responsibilities
Include 3-5 results-oriented bullet points per section
Bullet point should only be 1-2 lines long
Add measurable metrics and results to your bullets whenever possible

[DO NOT ADD ANY EXTRA SECTION OR HEADINGS OR SUB HEADINGS EXCEPT THE ONE SPECIFIED in ABOVE FORMAT. ONLY OUTPUT ABOVE SECTION NO EXTRA CONTENT OR SECTION REQUIRED. FOLLOW THE FORMAT STRICTLY]
Resume Text:
'''

EDUCATION_SECTION = '''
## EDUCATION
[OUTPUT ONLY THE FIRST 3 EDUCATION OUT OF THE TEXT IGNORE ALL REST, FOLLOW THE EDUCATION format specifically, write in short. Do not write more than 3 EDUCATION, write only first 3. Thier content should follow below format]

**Write here the extracted Title of education**
<span>05/2017 - 01/2019</span> [WRITE THE DURATION PERIOD ENCLOSED WITH <span> tag]
(Write the Program/Specialty)
Write here the extracted University/Institution Name, Location, Honors

[DO NOT ADD ANY EXTRA SECTION OR HEADINGS OR SUB HEADINGS EXCEPT THE ONE SPECIFIED in ABOVE FORMAT. ONLY OUTPUT ABOVE SECTION NO EXTRA CONTENT OR SECTION REQUIRED. FOLLOW THE FORMAT STRICTLY]
Resume Text:
'''

ADDITIONAL_INFO_SECTION = '''
## ADDITIONAL INFORMATION
[Write the content only in one line short, should not exceed one line length, DO NOT WRITE ANY EXTRA POINTS except the "Technical Skills", 
"Certifications", "Awards", "Languages", "Interests" if any of these is not given SIMPLY IGNORE IT AND DO NOT WRITE JUST THE HEADING, ONLY write with the actual data if present.
 IGNORE the REST of the content. OUTPUT SHOULD ONLY BE IN FORMAT IF something is not given just ignore it]

[SHOULD BE ONLY A MARKDOWN List in order of relevancy to job application.]
  + **Technical Skills**: (Just output level of proficiency) i.e. SQL (advanced), HTML (proficient), etc.
  + **Certifications**: i.e. Black Belt in Six Sigma, PMP
  + **Awards**: i.e. Officer of the Year (2014), ranked #1 of 80 for enhancements to operational efficiency
  + **Languages**: (Just output level of proficiency), i.e. French (advanced), Spanish (basic), German (conversational)
  + **Interests**: (include examples), i.e. running (completed 4 marathons across 3 continents), cooking (won Ft. Benning’s annual Chili Cook-Off)

[DO NOT ADD ANY EXTRA SECTION OR HEADINGS OR SUB HEADINGS EXCEPT THE ONE SPECIFIED in ABOVE FORMAT. ONLY OUTPUT ABOVE SECTION NO EXTRA CONTENT OR SECTION REQUIRED. FOLLOW THE FORMAT STRICTLY]
Resume text:
'''


from .chat_ml import build_embeddings, search_query
import re 

def clean_resume(markdown_text):
  desired_sections = ["## SUMMARY", "## EXPERIENCE", "## EDUCATION", "## ADDITIONAL INFORMATION"]

  # Find all the headings and their contents
  headings = re.findall(r'## (.+)', markdown_text)
  contents = re.split(r'## .+', markdown_text)[1:]


  # Filter out the desired sections
  filtered_headings = []
  filtered_contents = []

  for heading, content in zip(headings, contents):
      if str("## " + heading.strip()) in desired_sections:
          filtered_headings.append(heading)
          filtered_contents.append(content.strip())

  # Construct the cleaned markdown text
  cleaned_markdown_text = "\n\n".join([f"## {heading}\n\n{content}\n" for heading, content in zip(filtered_headings, filtered_contents)])
  return cleaned_markdown_text


def generate_pdf(text, _file_path, out_file):
  with open('style.css', 'r') as f:
    _STYLES = f.read()
  chunks = textwrap.wrap(text, 2000)

  FINAL_RESUME_CONTENT = ""
  FRONT_MATTER = generate_gpt(BASE_PROMPT + INFO_SECTION + chunks[0] + extract_links_from_pdf(_file_path), history=[]) + "\n\n"

  result = search_query('education experience achievements certificates rewards courses',  _file_path, text)
  _PTR = result
  if len(_PTR) > 9000:
    _PTR = _PTR[:9000]

  FINAL_RESUME_CONTENT += generate_gpt(BASE_PROMPT + SUMMARY_SECTION + _PTR, history=[]) + "\n\n"

  result = search_query('experience experiences',  _file_path, text)
  FINAL_RESUME_CONTENT += generate_gpt(BASE_PROMPT + EXPERIENCE_SECTION + result, history=[]) + "\n\n"

  result = search_query('education educational qualifications',  _file_path, text)
  FINAL_RESUME_CONTENT += generate_gpt(BASE_PROMPT + EDUCATION_SECTION + result, history=[]) + "\n\n"

  result = search_query('additional information achievements certificates rewards courses',  _file_path, text)
  FINAL_RESUME_CONTENT += generate_gpt(BASE_PROMPT + ADDITIONAL_INFO_SECTION + result, history=[]) + "\n\n"

  FINAL_RESUME_CONTENT_ = clean_resume(FINAL_RESUME_CONTENT) 
  FRONT_MATTER += FINAL_RESUME_CONTENT_

  print('--------------------------------------------------------------------------------------')
  print(FRONT_MATTER)

  _HTML = markdown.markdown(FRONT_MATTER)

  BASE_HTML_TEMPLATE = '''
<!doctype html>
<html>
<head>
<meta charset="utf-8"/>
<title>Marked in the browser</title>
<style>
{_STYLES}
</style>
</head>
<body>
<img class="_logo" src="https://dev-to-uploads.s3.amazonaws.com/uploads/articles/kokgebg96tn754ti1isq.png" alt="">
<div style="clear: both;"></div>
<div id="content" style="margin-top: -55px;">
{_HTML}
</div>
<script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
</body>
</html>
'''

  html_file = out_file + '.html' 
  with open(html_file, 'w', encoding='utf-8') as f:
    f.write(BASE_HTML_TEMPLATE.format(_STYLES=_STYLES, _HTML=_HTML).replace("|", "&#x7C;"))

  pdf_file = out_file + '.pdf'
  # pdf_file = 'output.pdf'
  # options = {
  #     'page-size': 'A4',
  #     'margin-top': '0mm',
  #     'margin-right': '0mm',
  #     'margin-bottom': '0mm',
  #     'margin-left': '0mm',
  # }
  # pdfkit.from_file(html_file, pdf_file, options=options)

  # ============= Using CONVERTAPI ===========
  import convertapi
  convertapi.api_secret = 'yRPEYmWFycEVpXiV'
  convertapi.convert('pdf', {
    'File': html_file
}, from_format = 'html').save_files(pdf_file)
  print(f'PDF file generated successfully at {pdf_file}')


import time     
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def modify_docx(file_path):
    doc = docx.Document(file_path)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if 'graphicData' in run._r.xml:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                break

        if "$$$" in paragraph.text:
            runs = paragraph.runs 
            print(runs[1].text)
            runs[0].bold = True
            runs[-2].text = runs[-2].text.replace("$$$", "")
            tab_stops = paragraph.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(2), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            tab_stops.add_tab_stop(Inches(7), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            new_text = f"{''.join([run.text for run in runs[1:-1]])}\t{runs[-1].text}"
            paragraph.clear()
            paragraph.add_run(runs[0].text).bold = True
            paragraph.add_run(new_text)

    doc.save('kkk' + file_path)



def create_access_token():
    url = "https://pdf-services-ue1.adobe.io/token"
    client_id = CLIENT_ID
    client_secret = CLIENT_SECRET

    payload = {
        "client_id": client_id,
        "client_secret": client_secret
    }

    headers = {
        "Content-Type": "application/x-www-form-urlencoded"
    }

    response = requests.post(url, data=payload, headers=headers)

    if response.status_code == 200:
        access_token = response.json().get("access_token")
        print("Access Token:", access_token)
        return access_token
    else:
        return None
    

def upload_pdf_file(upload_uri, file_path):
    media_type = 'application/pdf'
    with open(file_path, "rb") as file:
        files = {"file": (file.name, file, media_type)}
        headers = {
            "Content-Type": "application/pdf",
        }
        response = requests.put(upload_uri, files=files, headers=headers)
        print(response)
        print(response.text)


def get_upload_presigned_uri(access_token, file_path):
    url = "https://pdf-services-ue1.adobe.io/assets"  

    headers = {
        "Authorization": "Bearer " + access_token,
        "x-api-key": "71d2c5ded07748889b86e68730b5780e",
    }

    payload = {
        "mediaType": 'application/pdf'
    }

    response = requests.post(url, headers=headers, json=payload)

    if response.status_code == 200:
        data = response.json()
        asset_id = data["assetID"]
        upload_uri = data["uploadUri"]
        media_type = "application/pdf"

        upload_pdf_file(upload_uri, file_path)
        return asset_id
    return None


def get_download_uri(access_token, assetId):
    url = f"https://pdf-services-ue1.adobe.io/assets/{assetId}"
    headers = {
        "Authorization": "Bearer " + access_token,
        "x-api-key": CLIENT_ID,
    }

    response = requests.get(url, headers=headers)
    print(response)
    print(response.text)
    if response.status_code == 200:
        data = response.json()
        download_uri = data.get("downloadUri")
        size = data.get("size")
        resource_type = data.get("type")
        print("Download URI:", download_uri)
        print("Size:", size)
        print("Resource Type:", resource_type)
        return download_uri
    else:
        return None


def poll_export_pdf_job_status(access_token, url):
    headers = {
        "Authorization": "Bearer " + access_token,
        "x-api-key": CLIENT_ID
    }

    while True:
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            data = response.json()
            status = data.get("status")
            print("Job status:", status)

            if status == "done":
                asset_id = data['asset']['assetID']
                print(asset_id)
                return asset_id
            elif status == "failed":
                break
            else:
                time.sleep(5)
        else:
            print("Failed to poll job status.")
            break


def export_pdf_to_docx(access_token, asset_id):
    url = "https://pdf-services-ue1.adobe.io/operation/exportpdf"
    headers = {
        "Authorization": "Bearer " + access_token,
        "x-api-key": CLIENT_ID
    }
    payload = {
        "assetID": asset_id,
        "targetFormat": "docx"
    }

    response = requests.post(url, headers=headers, json=payload)
    print(response)

    if response.status_code == 201:
        location = response.headers.get("location")
        print("Export job status URI:", location)

        status = poll_export_pdf_job_status(access_token, location)
        return status
    else:
        return None
    

def download_file(url, file_path):
    response = requests.get(url)
    if response.status_code == 200:
        with open(file_path, 'wb') as file:
            file.write(response.content)
        print("File downloaded successfully.")
    else:
        print("Failed to download the file.")


from time import sleep

def convert_pdf_to_docx(file_path):
  access_token = create_access_token()
  # file_path = "output_result.pdf"

  out_url_id = get_upload_presigned_uri(access_token, file_path)
  print(out_url_id)

  down_url = get_download_uri(access_token, out_url_id)
  print(down_url)

  exported_doc_id = export_pdf_to_docx(access_token, out_url_id)
  print(exported_doc_id)
  sleep(3)

  exported_doc_url = get_download_uri(access_token, str(exported_doc_id))
  print(exported_doc_url)

  download_file(exported_doc_url, 'temp.docx')



# modify_docx('converted.docx')

# txt = extract_text_from_pdf('/content/resume_2.pdf')
# build_embeddings('/content/resume_2.pdf')
# generate_pdf(txt, '/content/resume_2.pdf')