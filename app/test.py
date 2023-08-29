  

# import markdown
# import json
# import os
# import requests
# import pdfkit
# import PyPDF2
# import textwrap



# # Path to the HTML file
html_file = 'page.html'
pdf_file = 'sksksks434.pdf'
# options = {
#     'page-size': 'A4',
#     'margin-top': '20mm',
#     'margin-right': '0mm',
#     'margin-bottom': '0mm',
#     'margin-left': '0mm',
# }
# pdfkit.from_file(html_file, pdf_file, options=options)




# from pdf2docx import parse

# pdf_file = pdf_file
docx_file = 'sample.docx'

# # convert pdf to docx
# parse(pdf_file, docx_file)


from htmldocx import HtmlToDocx

new_parser = HtmlToDocx()
new_parser.parse_html_file(html_file, docx_file)

# import convertapi
# convertapi.api_secret = 'yRPEYmWFycEVpXiV'
# convertapi.convert('docx', {
#     'File': html_file
# }, from_format = 'html').save_files(docx_file)



import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def modify_docx(file_path):
    doc = docx.Document(file_path)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if 'graphicData' in run._r.xml:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                break

        if "$$$" in paragraph.text:
            runs = paragraph.runs 
            runs[0].bold = True
            runs[-2].text = runs[-2].text.replace("$$$", "")
            tab_stops = paragraph.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(2), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            tab_stops.add_tab_stop(Inches(7), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
            new_text = f"{''.join([run.text for run in runs[1:-1]])}\t{runs[-1].text}"
            paragraph.clear()
            paragraph.add_run(runs[0].text).bold = True
            paragraph.add_run(new_text)

    doc.save('dddd' + file_path)

# Usage
file_path = 'sample.docx'
modify_docx(file_path)




